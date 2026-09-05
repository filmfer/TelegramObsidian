"""Document text extraction: PDF, DOCX, XLSX, TXT, JSON, MD, CSV, EML.

Every extractor is *pure* — it takes a filesystem path and returns plain
text — so parsers can be unit-tested without Telegram or network access.
`extract_document_text()` is the single entry point bot.py calls; it
dispatches on file extension and raises ValueError for unsupported types.
"""
from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Supported file extensions for text extraction
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".json", ".xlsx", ".csv", ".eml"}


def parse_document(file_path: str) -> str | None:
    """Extract plain text from a document based on extension. Returns None if unsupported."""
    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        logger.warning(f"Unsupported file type: {ext}")
        return None

    dispatch = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".eml": _extract_eml,
    }
    handler = dispatch.get(ext)
    if handler:
        return handler(file_path)
    return _extract_text(file_path)


def _extract_pdf(file_path: str) -> str | None:
    """Extract text from a PDF using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text_pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(text_pages).strip()
    except Exception as e:
        logger.error(f"Failed to extract PDF text: {e}")
        return None


def _extract_docx(file_path: str) -> str | None:
    """Extract text from DOCX using python-docx."""
    try:
        import docx

        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"Failed to extract DOCX text: {e}")
        return None


def _extract_xlsx(file_path: str) -> str | None:
    """Extract text from XLSX via openpyxl."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"### Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    lines.append(row_text)
        return "\n".join(lines).strip()
    except Exception as e:
        logger.error(f"Failed to extract XLSX text: {e}")
        return None


def _extract_text(file_path: str) -> str | None:
    """Read plain-text files: .txt, .md, .json, .csv."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        logger.error(f"Failed to read text file: {e}")
        return None


def _extract_eml(file_path: str) -> str | None:
    """Extract plain text from an EML email file."""
    try:
        from email import message_from_binary_file
        from bs4 import BeautifulSoup

        with open(file_path, "rb") as f:
            msg = message_from_binary_file(f)
        pieces = []
        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                payload = part.get_payload(decode=True)
                if not payload:
                    continue
                charset = part.get_content_charset() or "utf-8"
                if ctype == "text/plain":
                    pieces.append(payload.decode(charset, errors="ignore"))
                elif ctype == "text/html":
                    pieces.append(BeautifulSoup(payload, "html.parser").get_text())
        else:
            payload = msg.get_payload(decode=True)
            if not payload:
                return None
            charset = msg.get_content_charset() or "utf-8"
            if msg.get_content_type() == "text/html":
                pieces.append(BeautifulSoup(payload, "html.parser").get_text())
            else:
                pieces.append(payload.decode(charset, errors="ignore"))
        result = "\n".join(pieces).strip()
        return result or None
    except Exception as e:
        logger.error(f"Failed to parse EML: {e}")
        return None
